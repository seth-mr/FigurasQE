from __future__ import annotations

import argparse
from html import escape
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
HEADER_FILL = "F2F4F7"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def configure_document(title: str, subtitle: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("FIGURAS QUE ENSEÑAN  |  PROYECTO FINAL")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Equipo FigurasQE")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(14)
    title_p.paragraph_format.space_after = Pt(5)
    title_run = title_p.add_run(title)
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = DARK_BLUE

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(16)
    subtitle_run = subtitle_p.add_run(subtitle)
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = MUTED

    metadata = doc.add_table(rows=3, cols=2)
    metadata.style = "Table Grid"
    set_table_geometry(metadata, [2700, 6660])
    values = (
        ("Proyecto", "Figuras Que Enseñan"),
        ("Asignatura", "Desarrollo de Sistemas en Red"),
        ("Fecha", "12 de junio de 2026"),
    )
    for row, (label, value) in zip(metadata.rows, values):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], HEADER_FILL)
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    return doc


def add_table(doc, headers: list[str], rows: list[tuple[str, ...]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, HEADER_FILL)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
        set_table_geometry(table, widths)
    return table


def iter_blocks(doc: Document):
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield DocxParagraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, doc)


def pdf_styles():
    styles = getSampleStyleSheet()
    return {
        "Normal": ParagraphStyle(
            "FQENormal",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=13,
            textColor=colors.HexColor("#202124"),
            spaceAfter=6,
        ),
        "Heading 1": ParagraphStyle(
            "FQEH1",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=19,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=16,
            spaceAfter=8,
            keepWithNext=True,
        ),
        "Heading 2": ParagraphStyle(
            "FQEH2",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=16,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=12,
            spaceAfter=6,
            keepWithNext=True,
        ),
        "Title": ParagraphStyle(
            "FQETitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=23,
            leading=27,
            textColor=colors.HexColor("#1F4D78"),
            alignment=TA_LEFT,
            spaceAfter=5,
        ),
        "Subtitle": ParagraphStyle(
            "FQESubtitle",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=12,
            leading=15,
            textColor=colors.HexColor("#5A5A5A"),
            spaceAfter=14,
        ),
        "List": ParagraphStyle(
            "FQEList",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=13,
            leftIndent=0.28 * inch,
            firstLineIndent=-0.18 * inch,
            spaceAfter=5,
        ),
    }


def export_pdf(docx_path: Path, pdf_path: Path) -> None:
    source = Document(docx_path)
    styles = pdf_styles()
    story = []
    list_number = 0
    paragraph_index = 0

    for block in iter_blocks(source):
        if isinstance(block, DocxParagraph):
            text = block.text.strip()
            if not text:
                story.append(Spacer(1, 5))
                continue

            style_name = block.style.name if block.style else "Normal"
            if paragraph_index == 0:
                style = styles["Title"]
            elif paragraph_index == 1:
                style = styles["Subtitle"]
            elif style_name in ("Heading 1", "Heading 2"):
                style = styles[style_name]
                list_number = 0
            elif style_name == "List Bullet":
                style = styles["List"]
                text = f"• {text}"
            elif style_name == "List Number":
                style = styles["List"]
                list_number += 1
                text = f"{list_number}. {text}"
            else:
                style = styles["Normal"]

            story.append(Paragraph(escape(text), style))
            paragraph_index += 1
            continue

        rows = []
        for row in block.rows:
            rows.append(
                [Paragraph(escape(cell.text.strip()), styles["Normal"]) for cell in row.cells]
            )
        columns = len(rows[0]) if rows else 0
        if columns == 2:
            widths = [1.75 * inch, 4.75 * inch]
        elif columns == 3:
            widths = [0.8 * inch, 1.15 * inch, 4.55 * inch]
        elif columns == 4:
            widths = [0.65 * inch, 1.35 * inch, 1.35 * inch, 3.15 * inch]
        else:
            widths = [6.5 * inch / max(columns, 1)] * max(columns, 1)

        table = Table(rows, colWidths=widths, repeatRows=1, hAlign="LEFT")
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#202124")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B8BEC6")),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]
            )
        )
        story.extend((table, Spacer(1, 7)))

    def draw_page(canvas, _doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8.5)
        canvas.setFillColor(colors.HexColor("#666666"))
        canvas.drawString(inch, 10.55 * inch, "FIGURAS QUE ENSEÑAN  |  PROYECTO FINAL")
        canvas.drawRightString(7.5 * inch, 0.48 * inch, f"Equipo FigurasQE  |  Página {_doc.page}")
        canvas.restoreState()

    pdf = BaseDocTemplate(
        str(pdf_path),
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=0.72 * inch,
        bottomMargin=0.72 * inch,
        title=pdf_path.stem,
        author="Equipo FigurasQE",
    )
    frame = Frame(pdf.leftMargin, pdf.bottomMargin, pdf.width, pdf.height, id="main")
    pdf.addPageTemplates(PageTemplate(id="standard", frames=[frame], onPage=draw_page))
    pdf.build(story)


def build_problem_definition(output: Path) -> None:
    doc = configure_document(
        "AR01 - Definición del Problema",
        "Versión actualizada del planteamiento, alcance y solución propuesta",
    )

    doc.add_heading("1. Contexto", level=1)
    doc.add_paragraph(
        "El aprendizaje inicial de la suma suele apoyarse en ejercicios repetitivos "
        "que no siempre conectan la representación numérica con una experiencia "
        "concreta. Para niñas y niños con habilidades lectoras en desarrollo, una "
        "interfaz saturada o poco interactiva puede reducir la motivación y dificultar "
        "la comprensión de las cantidades."
    )
    doc.add_paragraph(
        "Al mismo tiempo, tutores, profesores y responsables necesitan información "
        "objetiva sobre sesiones, resultados e intentos para acompañar el progreso. "
        "Cuando esa información queda dispersa o sólo se observa durante la actividad, "
        "el seguimiento depende de la presencia constante del adulto."
    )

    doc.add_heading("2. Problema identificado", level=1)
    doc.add_paragraph(
        "Se requiere una plataforma educativa que combine interacción corporal, "
        "retroalimentación visual y seguimiento del aprendizaje en un entorno que "
        "pueda utilizarse desde distintos clientes. Además, la plataforma debe "
        "centralizar autenticación, datos, logs y monitoreo sin exponer directamente "
        "los servicios internos."
    )
    for item in (
        "Falta de una experiencia lúdica que vincule dedos, cantidades y operaciones.",
        "Dificultad para registrar sesiones y resultados de manera consistente.",
        "Ausencia de una vista unificada del progreso para estudiantes y tutores.",
        "Necesidad de separar roles y proteger la información de cada usuario.",
        "Necesidad de operar desde varios equipos con acceso seguro a la cámara.",
    ):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. Usuarios afectados", level=1)
    add_table(
        doc,
        ["Actor", "Necesidad principal"],
        [
            ("Estudiante", "Aprender y practicar sumas mediante niveles claros, visuales e interactivos."),
            ("Tutor, profesor o padre", "Asociar estudiantes y consultar sesiones, resultados y progreso."),
            ("Administrador", "Revisar estadísticas, logs, salud de servicios y operación general."),
        ],
        [2500, 6860],
    )

    doc.add_heading("4. Objetivo general", level=1)
    doc.add_paragraph(
        "Desarrollar un sistema distribuido de enseñanza que permita practicar sumas "
        "mediante actividades visuales y detección de dedos, registre el avance y "
        "ofrezca seguimiento por roles desde clientes web, móvil y administrativo."
    )

    doc.add_heading("5. Objetivos específicos", level=1)
    for item in (
        "Proporcionar niveles educativos con retroalimentación inmediata.",
        "Procesar imágenes de manos y devolver el conteo de dedos por cliente.",
        "Registrar usuarios, sesiones, niveles y resultados en PostgreSQL.",
        "Permitir a tutores consultar únicamente a sus estudiantes asociados.",
        "Centralizar solicitudes mediante un API Gateway.",
        "Registrar eventos técnicos con RabbitMQ y MongoDB.",
        "Ofrecer HTTPS para que varios navegadores puedan utilizar la cámara.",
    ):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("6. Alcance", level=1)
    doc.add_paragraph(
        "El alcance incluye clientes web, Android y WPF; autenticación JWT; gestión de "
        "estudiantes, tutores, sesiones, niveles y resultados; detección de manos; "
        "mensajería; almacenamiento de logs; monitoreo; contenedores Docker y acceso "
        "HTTPS para pruebas locales y remotas."
    )
    doc.add_paragraph(
        "El proyecto no plantea un diagnóstico pedagógico o médico, ni sustituye la "
        "evaluación profesional. El túnel público incluido es temporal y está destinado "
        "a pruebas, no a un despliegue productivo permanente."
    )

    doc.add_heading("7. Solución propuesta", level=1)
    doc.add_paragraph(
        "Figuras Que Enseñan adopta una arquitectura modular. Caddy ofrece HTTPS; el "
        "frontend consume rutas relativas; el Gateway centraliza las operaciones y "
        "traduce la detección de manos hacia gRPC. Los servicios de autenticación y "
        "datos usan PostgreSQL, mientras que RabbitMQ, un listener y MongoDB manejan "
        "eventos operativos. El servicio de manos mantiene detectores independientes "
        "para atender varios clientes sin compartir estado de seguimiento."
    )

    doc.add_heading("8. Criterios de éxito", level=1)
    for item in (
        "Un estudiante puede autenticarse, jugar y registrar su resultado.",
        "Dos o más clientes pueden detectar manos simultáneamente sin bloquearse.",
        "Un tutor puede consultar sólo el progreso de estudiantes asociados.",
        "La cámara funciona desde otros equipos mediante un origen HTTPS confiable.",
        "Los servicios pueden desplegarse y verificarse con Docker Compose.",
    ):
        doc.add_paragraph(item, style="List Bullet")

    doc.save(output)


def build_functional_spec(output: Path) -> None:
    doc = configure_document(
        "AR02 - Especificación Funcional",
        "Actores, capacidades, reglas de negocio y requisitos del sistema",
    )

    doc.add_heading("1. Propósito", level=1)
    doc.add_paragraph(
        "Esta especificación define el comportamiento funcional global de Figuras Que "
        "Enseñan y las restricciones principales que coordinan clientes, Gateway, "
        "servicios, persistencia y detección de manos."
    )

    doc.add_heading("2. Actores", level=1)
    add_table(
        doc,
        ["Actor", "Responsabilidades"],
        [
            ("Estudiante", "Registrarse, iniciar sesión, jugar niveles y consultar su información, sesiones y resultados."),
            ("Tutor / profesor / padre", "Gestionar su perfil, asociar estudiantes y revisar su progreso."),
            ("Administrador", "Consultar estadísticas, logs, estado de servicios y su perfil administrativo."),
        ],
        [2500, 6860],
    )

    doc.add_heading("3. Módulos funcionales", level=1)
    add_table(
        doc,
        ["Módulo", "Responsabilidad"],
        [
            ("Frontend", "Presentar la experiencia educativa y capturar la interacción del estudiante."),
            ("Gateway", "Centralizar rutas, autorización, salud y comunicación con servicios internos."),
            ("Autenticación", "Validar credenciales y emitir tokens JWT con identidad y rol."),
            ("Datos", "Administrar estudiantes, tutores, niveles, sesiones y resultados."),
            ("HandsDetection", "Procesar imágenes y devolver conteos independientes por solicitud."),
            ("Logs y mensajería", "Publicar, persistir y consultar eventos operativos."),
        ],
        [2200, 7160],
    )

    doc.add_heading("4. Requisitos funcionales", level=1)
    functional = [
        ("RF-01", "Alta", "Registrar usuarios con rol de estudiante o tutor."),
        ("RF-02", "Alta", "Autenticar estudiantes, tutores y administradores."),
        ("RF-03", "Alta", "Emitir y validar tokens JWT con identidad y rol."),
        ("RF-04", "Alta", "Dirigir a cada usuario a las funciones permitidas por su rol."),
        ("RF-05", "Alta", "Consultar perfil, tutor, sesiones y resultados del estudiante."),
        ("RF-06", "Alta", "Consultar y actualizar el perfil del tutor."),
        ("RF-07", "Alta", "Asociar un estudiante a un tutor mediante correo."),
        ("RF-08", "Alta", "Consultar estudiantes asignados y su progreso."),
        ("RF-09", "Alta", "Registrar sesiones con fechas y dispositivo."),
        ("RF-10", "Alta", "Registrar resultados con intentos, fallos, tiempo y estado."),
        ("RF-11", "Alta", "Enviar imágenes al Gateway y obtener el conteo de dedos."),
        ("RF-12", "Alta", "Atender detección de manos para varios clientes concurrentes."),
        ("RF-13", "Media", "Consultar estadísticas generales del sistema."),
        ("RF-14", "Media", "Publicar, persistir y consultar logs operativos."),
        ("RF-15", "Media", "Consultar endpoints de salud de los servicios."),
        ("RF-16", "Alta", "Ofrecer acceso HTTPS para habilitar la cámara del navegador."),
    ]
    add_table(doc, ["ID", "Prioridad", "Requisito"], functional, [1000, 1300, 7060])

    doc.add_heading("5. Requisitos no funcionales", level=1)
    non_functional = [
        ("RNF-01", "Seguridad", "Aplicar autenticación y autorización basada en roles."),
        ("RNF-02", "Privacidad", "Limitar los datos al propietario o tutor asociado."),
        ("RNF-03", "Configuración", "Usar variables de entorno para secretos, credenciales y capacidad."),
        ("RNF-04", "Arquitectura", "Mantener responsabilidades separadas por servicio."),
        ("RNF-05", "Interoperabilidad", "Comunicar clientes por HTTP/HTTPS y detección por gRPC."),
        ("RNF-06", "Usabilidad", "Presentar una interfaz clara para usuarios infantiles."),
        ("RNF-07", "Observabilidad", "Registrar eventos y exponer comprobaciones de salud."),
        ("RNF-08", "Despliegue", "Ejecutar el sistema completo mediante Docker Compose."),
        ("RNF-09", "Concurrencia", "Evitar estado global de seguimiento entre clientes de manos."),
        ("RNF-10", "Compatibilidad", "Utilizar un origen HTTPS reconocido para acceso a cámara."),
    ]
    add_table(doc, ["ID", "Categoría", "Requisito"], non_functional, [1100, 1700, 6560])

    doc.add_heading("6. Casos de uso principales", level=1)
    use_cases = [
        ("CU-01", "Iniciar sesión", "Todos", "El usuario recibe un JWT y accede según su rol."),
        ("CU-02", "Registrar usuario", "Estudiante / tutor", "Se valida el correo y se crea la cuenta."),
        ("CU-03", "Consultar progreso", "Estudiante / tutor", "Se devuelven sesiones y resultados autorizados."),
        ("CU-04", "Jugar nivel", "Estudiante", "Se procesa la interacción y se registra el resultado."),
        ("CU-05", "Detectar dedos", "Estudiante", "La imagen viaja por Gateway y gRPC hasta HandsDetection."),
        ("CU-06", "Consultar logs", "Administrador", "Se filtran eventos persistidos en MongoDB."),
        ("CU-07", "Monitorear servicios", "Administrador", "Se consulta la salud de componentes conocidos."),
    ]
    add_table(doc, ["ID", "Caso", "Actor", "Resultado"], use_cases, [900, 1900, 1900, 4660])

    doc.add_heading("7. Flujo funcional: jugar con detección de manos", level=1)
    steps = (
        "El estudiante inicia sesión y abre un nivel.",
        "El navegador solicita permiso de cámara desde un origen HTTPS.",
        "El frontend captura una imagen e incluye un identificador estable de la pestaña.",
        "El Gateway valida la solicitud y la envía al servicio de manos mediante gRPC.",
        "HandsDetection toma una instancia disponible de su grupo de detectores.",
        "El servicio devuelve mano izquierda, mano derecha, total y estado de detección.",
        "El frontend evalúa la respuesta y continúa el nivel.",
        "Al finalizar, el resultado se persiste mediante el servicio de datos.",
    )
    for step in steps:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("8. Reglas de negocio", level=1)
    for item in (
        "El estudiante sólo consulta información asociada a su identificador.",
        "El tutor sólo consulta estudiantes vinculados a su cuenta.",
        "Las funciones administrativas requieren rol de administrador.",
        "Las solicitudes protegidas deben incluir un JWT válido.",
        "El identificador de cliente de manos distingue pestañas y sesiones de navegador.",
        "Cada solicitud de detección usa un detector disponible y lo libera al finalizar.",
    ):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("9. Manejo de errores", level=1)
    add_table(
        doc,
        ["Condición", "Respuesta esperada"],
        [
            ("Credenciales inválidas", "Rechazar el acceso sin emitir token."),
            ("JWT ausente o vencido", "Responder con estado de no autorizado."),
            ("Datos fuera del alcance del rol", "Negar el acceso."),
            ("Servicio interno no disponible", "Devolver un error controlado desde el Gateway."),
            ("Cámara bloqueada", "Informar que se requiere HTTPS y permiso del navegador."),
            ("No se detectan manos", "Devolver conteos en cero sin afectar otros clientes."),
        ],
        [3300, 6060],
    )

    doc.add_heading("10. Criterios de aceptación", level=1)
    for item in (
        "El stack inicia con un único comando y publica una URL HTTPS.",
        "El inicio de sesión respeta roles y permisos.",
        "Las sesiones y resultados quedan disponibles para consultas posteriores.",
        "Dos clientes simultáneos reciben detecciones independientes.",
        "Los logs y endpoints de salud permiten verificar la operación.",
    ):
        doc.add_paragraph(item, style="List Bullet")

    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("output_dir", type=Path)
    parser.add_argument("--team", default="FigurasQE")
    args = parser.parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)
    ar01_docx = args.output_dir / f"AR01_{args.team}.docx"
    ar02_docx = args.output_dir / f"AR02_{args.team}.docx"
    build_problem_definition(ar01_docx)
    build_functional_spec(ar02_docx)
    export_pdf(ar01_docx, args.output_dir / f"AR01_{args.team}.pdf")
    export_pdf(ar02_docx, args.output_dir / f"AR02_{args.team}.pdf")


if __name__ == "__main__":
    main()
